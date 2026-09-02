from pathlib import Path
from docx import Document

source = Path("/Users/thanhlt/Downloads/Kiến trúc hệ thống quản lý đội xe vận tải_.docx")
doc = Document(source)
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f"P{i:03d} [{p.style.name}] {text}")
for ti, table in enumerate(doc.tables):
    print(f"\nTABLE {ti}")
    for row in table.rows:
        print(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells))
