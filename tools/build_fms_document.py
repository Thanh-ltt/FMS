from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path("FMS_Kien_truc_he_thong_hoan_chinh.docx")
BLUE = "1F4E78"
DARK = "17324D"
LIGHT = "E8F0F7"
PALE = "F4F7FA"
GRAY = "5B6573"
WHITE = "FFFFFF"
BLACK = "1F1F1F"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.78)
sec.left_margin = sec.right_margin = Inches(0.85)
sec.header_distance = sec.footer_distance = Inches(0.35)

def font(run, name="Arial", size=10.5, bold=False, color=BLACK, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold, run.italic = bold, italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run

styles = doc.styles
normal = styles["Normal"]
normal.font.name, normal.font.size = "Arial", Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.15
for name, size, before, after, color in [
    ("Heading 1", 16, 14, 7, BLUE),
    ("Heading 2", 12.5, 10, 5, BLUE),
    ("Heading 3", 11, 7, 3, DARK),
]:
    s = styles[name]
    s.font.name, s.font.size, s.font.bold = "Arial", Pt(size), True
    s.font.color.rgb = RGBColor.from_string(color)
    s._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(before), Pt(after)
    s.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number", "List Number 2"):
    s = styles[name]
    s.font.name, s.font.size = "Arial", Pt(10.5)
    s._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    s.paragraph_format.left_indent = Inches(0.38)
    s.paragraph_format.first_line_indent = Inches(-0.19)
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.line_spacing = 1.15

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    rep = OxmlElement("w:tblHeader")
    rep.set(qn("w:val"), "true")
    trPr.append(rep)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None: tcPr.append(shd)

def cell_margin(cell, top=80, start=110, bottom=80, end=110):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if tcMar.getparent() is None: tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        x = tcMar.find(qn(f"w:{m}")) or OxmlElement(f"w:{m}")
        x.set(qn("w:w"), str(v)); x.set(qn("w:type"), "dxa")
        if x.getparent() is None: tcMar.append(x)

def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        c = OxmlElement("w:gridCol"); c.set(qn("w:w"), str(w)); grid.append(c)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tcW = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcW.set(qn("w:w"), str(widths[i])); tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margin(cell)

def add_table(headers, rows, widths, small=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_geometry(t, widths)
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        shade(hdr.cells[i], BLUE)
        p = hdr.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(text), size=small, bold=True, color=WHITE)
    for ri, values in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(values):
            if ri % 2: shade(cells[i], PALE)
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
            font(p.add_run(str(value)), size=small, color=BLACK)
    set_table_geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t

def para(text="", bold_lead=None, align=None, italic=False, color=BLACK, after=5):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), bold=True, color=color)
        font(p.add_run(text[len(bold_lead):]), color=color, italic=italic)
    else: font(p.add_run(text), color=color, italic=italic)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); font(p.add_run(text)); return p

def number(text, style="List Number"):
    p = doc.add_paragraph(style=style); font(p.add_run(text)); return p

def heading(text, level=1):
    return doc.add_heading(text, level=level)

def callout(label, text):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; set_table_geometry(t, [9120])
    c = t.cell(0,0); shade(c, LIGHT); cell_margin(c, 140, 160, 140, 160)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    font(p.add_run(label + "  "), bold=True, color=BLUE)
    font(p.add_run(text), color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(paragraph.add_run("FMS • Tài liệu kiến trúc  |  "), size=8.5, color=GRAY)
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(header.add_run("FLEET MANAGEMENT SYSTEM"), size=8, bold=True, color=GRAY)
add_page_number(sec.footer.paragraphs[0])

# Cover
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(115); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("TÀI LIỆU KIẾN TRÚC HỆ THỐNG"), size=12, bold=True, color=BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(8)
font(p.add_run("HỆ THỐNG QUẢN LÝ ĐỘI XE VẬN TẢI"), size=27, bold=True, color=DARK)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(32)
font(p.add_run("Fleet Management System (FMS)"), size=15, color=BLUE)
callout("PHẠM VI", "Kiến trúc, mô-đun nghiệp vụ, dữ liệu, API, bảo mật và hướng triển khai — được đối chiếu với mã nguồn hiện tại.")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(80)
font(p.add_run("Phiên bản tài liệu 1.0  •  07/2026"), size=10, color=GRAY)
doc.add_page_break()

heading("Mục lục", 1)
add_table(["Phần 1", "Phần 2"], [
    ("1. Tổng quan hệ thống", "6. Luồng nghiệp vụ trọng yếu"),
    ("2. Tác nhân và quyền truy cập", "7. Thiết kế REST API"),
    ("3. Kiến trúc tổng thể", "8. Bảo mật"),
    ("4. Các mô-đun nghiệp vụ", "9. Triển khai và vận hành"),
    ("5. Mô hình dữ liệu", "10. Đánh giá và hướng phát triển"),
    ("Phụ lục A. Cấu trúc project", "Phụ lục B. Ghi chú đối chiếu"),
], [4560, 4560], small=9)
heading("Thông tin tài liệu", 2)
add_table(["Thuộc tính", "Nội dung"], [
    ("Mục đích", "Mô tả ngắn gọn nhưng đầy đủ kiến trúc và thiết kế triển khai của FMS."),
    ("Nguồn đối chiếu", "Mã nguồn backend đa module, frontend React, cấu hình ứng dụng và migration cơ sở dữ liệu."),
    ("Đối tượng đọc", "Nhóm phát triển, giảng viên/người đánh giá, quản trị hệ thống và bên vận hành."),
    ("Trạng thái", "Phản ánh hiện trạng project tại thời điểm biên soạn."),
], [1900, 7220])

heading("1. Tổng quan hệ thống", 1)
heading("1.1 Mục tiêu", 2)
para("FMS số hóa quy trình vận tải từ quản lý nguồn lực đến kết toán: tiếp nhận khách hàng, lập hợp đồng, thu tiền đặt cọc, điều phối xe và tài xế, theo dõi chuyến, ghi nhận chi phí, bảo trì phương tiện, xuất hóa đơn và tổng hợp báo cáo tài chính.")
heading("1.2 Phạm vi chức năng", 2)
for x in [
    "Quản trị người dùng, nhân viên và phân quyền theo vai trò.",
    "Quản lý phương tiện, tài xế, trạng thái sẵn sàng và hồ sơ liên quan.",
    "Quản lý khách hàng, tài khoản khách hàng và cổng tra cứu cá nhân.",
    "Quản lý hợp đồng, đơn giá hàng hóa, giá cước và khoản đặt cọc/hoàn cọc.",
    "Lập và điều phối chuyến; kiểm tra điều kiện khởi hành; theo dõi trạng thái chuyến.",
    "Ghi nhận, duyệt hoặc từ chối chi phí chuyến; quản lý bảo trì phương tiện.",
    "Lập hóa đơn, phân bổ tiền cọc, ghi nhận thanh toán, công nợ và doanh thu.",
    "Báo cáo tài chính tổng hợp cho quản lý và kế toán.",
]: bullet(x)
heading("1.3 Ngoài phạm vi hiện tại", 2)
para("Project chưa thể hiện tích hợp GPS thời gian thực, tối ưu tuyến bằng thuật toán, ứng dụng di động native, gửi thông báo qua SMS/email hay tích hợp cổng thanh toán. Các nội dung này nên được xem là hướng mở rộng, không phải chức năng đã hoàn thành.")

heading("2. Tác nhân và quyền truy cập", 1)
add_table(["Vai trò", "Trách nhiệm chính", "Phạm vi tiêu biểu"], [
    ("ADMIN", "Quản trị toàn hệ thống", "Người dùng/nhân viên; toàn bộ nghiệp vụ; quyền xóa dữ liệu."),
    ("MANAGER", "Điều hành vận tải", "Xe, tài xế, khách hàng, hợp đồng, chuyến, bảo trì, báo cáo."),
    ("ACCOUNTANT", "Tài chính – kế toán", "Hợp đồng (đọc), khách hàng, hóa đơn, cọc, chi phí, báo cáo."),
    ("DRIVER", "Thực hiện chuyến", "Hồ sơ cá nhân, chuyến được giao, kiểm tra sẵn sàng, chi phí chuyến."),
    ("CUSTOMER", "Theo dõi dịch vụ", "Cổng khách hàng: thông tin cá nhân và dữ liệu gắn với tài khoản."),
], [1200, 2750, 5170], small=8.7)
callout("NGUYÊN TẮC", "Frontend ẩn/hiện màn hình theo vai trò để cải thiện trải nghiệm; backend Spring Security mới là lớp thực thi quyền truy cập bắt buộc.")

heading("3. Kiến trúc tổng thể", 1)
heading("3.1 Mô hình", 2)
para("Hệ thống sử dụng kiến trúc phân lớp, tách frontend và backend qua REST API. Backend được tổ chức thành Maven multi-module nhằm cô lập trách nhiệm và kiểm soát chiều phụ thuộc.")
add_table(["Lớp", "Thành phần", "Trách nhiệm"], [
    ("Presentation", "React + Vite", "Giao diện, định tuyến, kiểm soát truy cập phía client, gọi API."),
    ("API/Controller", "controller, auth", "Nhận HTTP request, kiểm tra xác thực, ánh xạ DTO, trả response."),
    ("Business", "service", "Quy tắc nghiệp vụ, chuyển trạng thái, tính toán, điều phối repository."),
    ("Data access", "repository", "Spring Data JPA, truy vấn và lưu trữ entity."),
    ("Domain", "data, common", "Entity, enum, base model, validation, lỗi và response dùng chung."),
    ("Persistence", "PostgreSQL + Flyway", "Dữ liệu quan hệ và tiến hóa schema bằng migration."),
], [1450, 1800, 5870], small=8.8)
heading("3.2 Luồng xử lý yêu cầu", 2)
callout("LUỒNG CHÍNH", "React → Axios/JSON → Security filters (JWT, đổi mật khẩu bắt buộc) → Controller → Service → Repository/JPA → PostgreSQL → ApiResponse/JSON → React.")
heading("3.3 Cấu trúc module backend", 2)
add_table(["Module", "Nội dung chính"], [
    ("common", "BaseEntity, chuẩn response, validation pattern, mã lỗi và GlobalExceptionHandler."),
    ("data", "Entity, enum và converter của miền nghiệp vụ."),
    ("repository", "Các interface repository cho entity."),
    ("service", "Service interface/implementation, DTO, request và mapper."),
    ("auth", "Đăng nhập/đăng ký/đổi mật khẩu, JWT, bootstrap admin, security filters."),
    ("controller", "REST controller của các mô-đun nghiệp vụ."),
    ("api-web", "Điểm khởi chạy Spring Boot, application.yaml và Flyway migrations."),
], [1500, 7620], small=9)

heading("4. Các mô-đun nghiệp vụ", 1)
add_table(["Mô-đun", "Năng lực chính", "Dữ liệu trung tâm"], [
    ("Tài khoản & nhân sự", "Đăng nhập, đổi mật khẩu, nhân viên, kích hoạt tài khoản", "User, Role"),
    ("Đội xe", "CRUD xe, lọc trạng thái, tìm kiếm, kiểm tra tải trọng", "Vehicle"),
    ("Tài xế", "Hồ sơ, bằng lái, tài khoản tài xế, khả dụng", "Driver, User"),
    ("Khách hàng", "Hồ sơ, liên kết/tạo tài khoản, tra cứu hợp đồng và hóa đơn", "Customer, User"),
    ("Hợp đồng & giá", "Giá cước, loại hàng, giá trị hợp đồng, trạng thái", "Contract, CargoRate"),
    ("Tiền cọc", "Thu cọc, gắn hợp đồng/chuyến, hoàn cọc, tổng hợp", "Deposit, DepositRefund"),
    ("Chuyến vận tải", "Tạo/sửa, phân xe/tài xế, readiness, bắt đầu/hoàn tất/hủy", "Trip"),
    ("Chi phí", "Ghi nhận theo chuyến, duyệt/từ chối, tổng hợp", "Expense"),
    ("Bảo trì", "Lịch sử, chi phí, bắt đầu/hoàn tất/hủy bảo trì", "Maintenance"),
    ("Hóa đơn", "Tạo, áp dụng cọc, thanh toán, quá hạn, doanh thu/công nợ", "Invoice, Payment, Allocation"),
    ("Báo cáo", "Tổng hợp doanh thu, chi phí và kết quả tài chính", "Các DTO báo cáo"),
], [1550, 5050, 2520], small=8.2)

heading("5. Mô hình dữ liệu", 1)
heading("5.1 Thực thể và quan hệ cốt lõi", 2)
add_table(["Thực thể", "Quan hệ nổi bật", "Thuộc tính nghiệp vụ tiêu biểu"], [
    ("User", "1–1 Customer; liên kết nghiệp vụ với Driver", "username, role, hồ sơ nhân viên, active, mustChangePassword"),
    ("Customer", "1–N Contract; 1–N Trip; 1–N Invoice/Deposit", "thông tin liên hệ, tài khoản liên kết"),
    ("Contract", "N–1 Customer; 1–N Trip/Deposit", "mã, thời hạn, loại hàng, giá cước, giá trị, điều khoản cọc"),
    ("Trip", "N–1 Vehicle/Driver/Customer/Contract", "tuyến, thời gian, cự ly, tải hàng, giá cước, trạng thái"),
    ("Invoice", "N–1 Customer/Trip; 1–N Payment/Allocation", "tổng tiền, tiền cọc áp dụng, đã trả, hạn trả, trạng thái"),
    ("Deposit", "N–1 Customer; tùy chọn Contract/Trip", "mã cọc, số tiền, phạm vi, mục đích, trạng thái"),
    ("Expense", "N–1 Trip", "nhóm chi phí, số tiền, chứng từ/mô tả, trạng thái duyệt"),
    ("Maintenance", "N–1 Vehicle", "loại bảo trì, thời gian, chi phí, trạng thái"),
], [1350, 3200, 4570], small=8.25)
heading("5.2 Quy ước kỹ thuật", 2)
for x in [
    "Khóa chính của entity sử dụng UUID dạng String.",
    "Các trạng thái và loại nghiệp vụ được lưu bằng EnumType.STRING để dễ đọc và an toàn khi thay đổi thứ tự enum.",
    "BaseEntity cung cấp dữ liệu dùng chung (ví dụ thời điểm tạo/cập nhật tùy triển khai).",
    "Flyway là nguồn thay đổi schema duy nhất; Hibernate dùng ddl-auto: validate để phát hiện migration bị thiếu.",
]: bullet(x)
heading("5.3 Trạng thái chính", 2)
add_table(["Đối tượng", "Trạng thái trong project"], [
    ("Trip", "CREATED, ASSIGNED, IN_PROGRESS, COMPLETED, CANCELLED"),
    ("Vehicle", "AVAILABLE, IN_TRIP, MAINTENANCE, INACTIVE"),
    ("Contract", "DRAFT, ACTIVE, COMPLETED, CANCELLED"),
    ("Invoice", "PENDING, PAID, OVERDUE, CANCELLED"),
    ("Expense", "PENDING, APPROVED, REJECTED"),
    ("Maintenance", "PENDING, IN_PROGRESS, COMPLETED, CANCELLED"),
], [1900, 7220], small=9)

heading("6. Luồng nghiệp vụ trọng yếu", 1)
heading("6.1 Từ hợp đồng đến hoàn tất chuyến", 2)
for x in [
    "Tạo hoặc chọn khách hàng; liên kết tài khoản nếu cần sử dụng cổng khách hàng.",
    "Lập hợp đồng, khai báo loại hàng, giá cước/giá trị, thời hạn và điều khoản tiền cọc.",
    "Ghi nhận khoản cọc theo hợp đồng hoặc theo chuyến; theo dõi trạng thái sử dụng/hoàn trả.",
    "Tạo chuyến từ thông tin hợp đồng; chọn xe và tài xế phù hợp.",
    "Kiểm tra readiness: xe khả dụng, không vướng bảo trì, tài xế hợp lệ và tải trọng đáp ứng.",
    "Bắt đầu chuyến, tài xế ghi nhận chi phí phát sinh; cấp có thẩm quyền duyệt hoặc từ chối.",
    "Hoàn thành chuyến, lập báo cáo chuyến và tạo hóa đơn cho khách hàng.",
    "Áp dụng khoản cọc, ghi nhận các lần thanh toán; cập nhật công nợ và báo cáo tài chính.",
]: number(x, style="List Number 2")
heading("6.2 Quy tắc cần bảo toàn", 2)
for x in [
    "Biển số xe và username không trùng; dữ liệu đầu vào tuân theo Bean Validation và ValidationPatterns.",
    "Cự ly, tải hàng, đơn giá và giá trị cước phải là số dương.",
    "Xe/tài xế chỉ được phân công khi đáp ứng điều kiện sẵn sàng; xe bảo trì không được vận hành.",
    "Hóa đơn PAID hoặc CANCELLED có số tiền phải thu bằng 0; PENDING quá hạn được biểu diễn hiệu lực là OVERDUE.",
    "Xóa dữ liệu nghiệp vụ được giới hạn cho ADMIN; thao tác trạng thái đi qua endpoint chuyên biệt.",
]: bullet(x)

heading("7. Thiết kế REST API", 1)
para("Backend sử dụng JSON qua HTTP, endpoint đặt theo tài nguyên. API hiện không có tiền tố /api; frontend gọi trực tiếp các path như /vehicles, /trips và /invoices.")
add_table(["Nhóm", "Endpoint tiêu biểu", "Mục đích"], [
    ("Xác thực", "POST /auth/login; /auth/register; /auth/change-password", "Phiên JWT và mật khẩu."),
    ("Xe", "GET/POST /vehicles; PATCH /vehicles/{id}/status", "Quản lý và đổi trạng thái xe."),
    ("Tài xế", "GET /drivers/available; POST /drivers/{id}/account", "Khả dụng và cấp tài khoản."),
    ("Hợp đồng", "POST /contracts; PATCH /contracts/{id}/activate|complete|cancel", "Vòng đời hợp đồng."),
    ("Chuyến", "POST /trips; GET /trips/{id}/readiness; PATCH /start|complete|cancel", "Điều phối và trạng thái chuyến."),
    ("Chi phí", "POST /expenses; PATCH /expenses/{id}/approve|reject", "Ghi nhận và xét duyệt."),
    ("Cọc", "POST /deposits; POST /deposits/{id}/refunds", "Thu và hoàn cọc."),
    ("Hóa đơn", "POST /invoices; PATCH /invoices/{id}/pay|cancel", "Lập hóa đơn và thanh toán."),
    ("Báo cáo", "GET /reports/financial", "Tổng hợp tài chính."),
    ("Cổng KH", "GET /customer-portal/me", "Dữ liệu theo tài khoản khách hàng."),
], [1250, 4350, 3520], small=8.2)
callout("CHUẨN PHẢN HỒI", "Các controller trả ApiResponse<T>; lỗi nghiệp vụ được chuẩn hóa qua ErrorCode, AppException và GlobalExceptionHandler.")

heading("8. Bảo mật", 1)
heading("8.1 Cơ chế hiện có", 2)
for x in [
    "Spring Security ở chế độ stateless; JWT được kiểm tra trước UsernamePasswordAuthenticationFilter.",
    "Mật khẩu được mã hóa; tài khoản tài xế có thể bị buộc đổi mật khẩu ở lần đăng nhập đầu.",
    "RBAC được cấu hình theo HTTP method và endpoint cho 5 vai trò.",
    "CORS hiện chỉ cho phép frontend local tại localhost/127.0.0.1:5173.",
    "CSRF bị tắt phù hợp với REST stateless dùng Authorization header.",
]: bullet(x)
heading("8.2 Rủi ro cấu hình cần xử lý trước production", 2)
add_table(["Mức", "Vấn đề", "Khuyến nghị"], [
    ("Cao", "Mật khẩu PostgreSQL và khóa JWT đang viết trực tiếp trong application.yaml.", "Chuyển sang biến môi trường/secret manager; thay khóa đã lộ; không commit secret."),
    ("Đã xử lý", "Hibernate từng tự cập nhật schema song song với Flyway.", "Đã chuyển sang validate; database mới dùng baseline B17 và thay đổi sau đó dùng V18+."),
    ("Vừa", "show-sql và log security ở DEBUG.", "Tắt trên production để giảm lộ dữ liệu và dung lượng log."),
    ("Vừa", "CORS cố định cho môi trường local.", "Tách cấu hình theo profile và whitelist đúng domain triển khai."),
], [900, 3650, 4570], small=8.2)

heading("9. Triển khai và vận hành", 1)
heading("9.1 Yêu cầu môi trường", 2)
add_table(["Thành phần", "Yêu cầu"], [
    ("Backend", "Java 21, Maven wrapper, Spring Boot 3.3.4"),
    ("Database", "PostgreSQL; database fms_db; baseline B17 cho database mới; migration tiếp theo từ V18"),
    ("Frontend", "Node.js tương thích Vite 8; React 19; chạy dev mặc định cổng 5173"),
    ("API", "Spring Boot mặc định cổng 8080"),
], [1800, 7320])
heading("9.2 Trình tự chạy local", 2)
for x in [
    "Tạo PostgreSQL database và khai báo credential qua cấu hình môi trường.",
    "Chạy backend từ module api-web bằng Maven wrapper; Flyway áp dụng migration khi khởi động.",
    "Cài dependency frontend, chạy Vite dev server và kiểm tra base URL API.",
    "Đăng nhập bằng tài khoản bootstrap hợp lệ, đổi mật khẩu nếu hệ thống yêu cầu.",
]: number(x)
heading("9.3 Kiểm thử", 2)
para("Project có Spring Boot tests, kiểm thử authentication/security, dữ liệu SQL cho luồng hoàn chỉnh và các trường hợp biên. Trước bàn giao nên chạy toàn bộ Maven test, frontend lint/build và một vòng kiểm thử tích hợp trên database sạch.")

heading("10. Đánh giá và hướng phát triển", 1)
heading("10.1 Điểm mạnh hiện tại", 2)
for x in [
    "Phân tách module rõ ràng giữa domain, repository, service, bảo mật và controller.",
    "Nghiệp vụ bao phủ chuỗi vận tải và tài chính tương đối đầy đủ, gồm cả đặt cọc và thanh toán nhiều lần.",
    "RBAC backend chi tiết, có cổng riêng cho tài xế và khách hàng.",
    "Schema được tiến hóa bằng migration và có dữ liệu kiểm thử nghiệp vụ.",
]: bullet(x)
heading("10.2 Ưu tiên cải tiến", 2)
add_table(["Ưu tiên", "Hạng mục", "Kết quả mong đợi"], [
    ("P0", "Quản lý secret và profile cấu hình", "An toàn khi triển khai, không lộ credential."),
    ("P0", "Chuẩn hóa Flyway/DDL", "Schema tái lập được và không drift."),
    ("P1", "OpenAPI/Swagger và versioning /api/v1", "API dễ tích hợp, có hợp đồng rõ ràng."),
    ("P1", "Test service/controller và integration test", "Giảm lỗi khi đổi quy tắc trạng thái."),
    ("P1", "Audit log cho thanh toán, duyệt chi phí, đổi trạng thái", "Truy vết nghiệp vụ tài chính."),
    ("P2", "Docker/CI/CD, health check, metrics", "Triển khai lặp lại và giám sát tốt hơn."),
    ("P2", "GPS, thông báo, tối ưu tuyến", "Mở rộng năng lực vận hành thời gian thực."),
], [950, 3800, 4370], small=8.5)

heading("Phụ lục A. Cấu trúc project rút gọn", 1)
tree = """FMS/
├── common/       # base model, response, validation, exception
├── data/         # entity, enum, converter
├── repository/   # Spring Data JPA repositories
├── service/      # service, implementation, DTO, request, mapper
├── auth/         # JWT, Spring Security, authentication
├── controller/   # REST controllers
├── api-web/      # Spring Boot entry point, config, Flyway migration
├── frontend/     # React, Vite, pages, components, services
└── test-data/    # SQL và hướng dẫn kiểm thử"""
p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25); p.paragraph_format.space_after=Pt(10)
font(p.add_run(tree), name="Courier New", size=9, color=DARK)
heading("Phụ lục B. Ghi chú đối chiếu", 1)
para("Tài liệu này thay thế các giả định trong bản nháp bằng thông tin đọc trực tiếp từ cấu trúc module, entity, enum, controller, security config, frontend routes, application config và migration hiện có. Khi project thay đổi API hoặc schema, tài liệu cần được cập nhật cùng phiên bản mã nguồn.")

# Keep rows together where practical and set metadata.
doc.core_properties.title = "Kiến trúc hệ thống quản lý đội xe vận tải (FMS)"
doc.core_properties.subject = "Tài liệu kiến trúc và thiết kế hệ thống"
doc.core_properties.author = "FMS Project Team"
doc.core_properties.keywords = "FMS, fleet management, Spring Boot, React, PostgreSQL"
doc.save(OUT)
print(OUT.resolve())
